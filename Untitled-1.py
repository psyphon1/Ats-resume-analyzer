"""
ATS Resume Analyzer - Production Ready | Enhanced Version
"""

import streamlit as st
import json, tempfile, os, re, io, logging, sys, time
from datetime import datetime
from typing import List, Dict, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

# Windows-compatible logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('ats.log', mode='w', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

try:
    import fitz  # PyMuPDF
    import pandas as pd
    from rapidocr_onnxruntime import RapidOCR
    from docx import Document
    from groq import Groq
    import gspread
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain.docstore.document import Document as LangchainDocument
    from langchain.chains import ConversationalRetrievalChain
    from langchain_groq import ChatGroq
    from langchain.memory import ConversationBufferMemory
    from langchain.prompts import PromptTemplate
    logger.info("All dependencies loaded successfully")
except ImportError as e:
    st.error(f"Missing library: {e}\nRun: pip install -r requirements.txt")
    st.stop()

CONFIG = {
    "groq_model": "llama-3.1-8b-instant",
    "embedding_model": "sentence-transformers/all-MiniLM-L6-v2",
    "chunk_size": 4000,  
    "chunk_overlap": 0,
    "max_resume_chars": 5000,
    "config_file": "api_config.json",
    "max_tokens": 1500,
    "temperature": 0.05,
    "max_workers": 3,
    "retry_delay": 2,
    "max_retries": 3
}

GOOGLE_SCOPES = [
    'https://www.googleapis.com/auth/spreadsheets',
    'https://www.googleapis.com/auth/drive.readonly'
]

class DocumentExtractor:
    def __init__(self):
        self.ocr = RapidOCR()
    
    def extract_text(self, path: str, ftype: str) -> str:
        try:
            if ftype == 'pdf':
                text_content = ""
                with fitz.open(path) as doc:
                    for page in doc:
                        page_text = page.get_text("text").strip()
                        if len(page_text) > 50: 
                            text_content += page_text + "\n"
                        else:
                            try:
                                pix = page.get_pixmap()
                                img_data = pix.tobytes("png")
                                result, _ = self.ocr(img_data)
                                if result:
                                    text_content += "\n".join([line[1] for line in result]) + "\n"
                            except Exception as e:
                                logger.error(f"OCR failed on PDF page: {e}")
                return text_content.strip()
            elif ftype == 'docx':
                return "\n".join([p.text.strip() for p in Document(path).paragraphs if p.text.strip()])
            elif ftype in ['png', 'jpg', 'jpeg']:
                result, _ = self.ocr(path)
                return "\n".join([line[1] for line in result]) if result else ""
        except Exception as e:
            logger.error(f"Extract error: {e}")
        return ""

class GoogleDriveHandler:
    def __init__(self, extractor):
        self.extractor = extractor
        creds = None
        if os.path.exists('token.json'):
            creds = Credentials.from_authorized_user_file('token.json', GOOGLE_SCOPES)
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', GOOGLE_SCOPES)
                creds = flow.run_local_server(port=0)
            with open('token.json', 'w') as token:
                token.write(creds.to_json())
        self.gc = gspread.authorize(creds)
        self.drive = build('drive', 'v3', credentials=creds)
    
    def process_sheet(self, url: str):
        try:
            sheet = self.gc.open_by_url(url).get_worksheet(0)
            records = sheet.get_all_records()
            
            link_col = next((k for r in records[:5] for k, v in r.items() if isinstance(v, str) and "drive.google.com" in v), None)
            
            if not link_col:
                yield {"error": "No Drive links found"}, 0, 0
                return
            
            for idx, row in enumerate(records, 1):
                try:
                    link = row.get(link_col, "").strip()
                    if not link:
                        continue
                    
                    match = re.search(r'/file/d/([a-zA-Z0-9_-]+)|id=([a-zA-Z0-9_-]+)', link)
                    if not match:
                        continue
                    fid = match.group(1) or match.group(2)
                    
                    meta = self.drive.files().get(fileId=fid).execute()
                    fname = f"temp_{fid}_{meta.get('name', 'file')}"
                    
                    request = self.drive.files().get_media(fileId=fid)
                    fh = io.BytesIO()
                    downloader = MediaIoBaseDownload(fh, request)
                    while not downloader.next_chunk()[1]:
                        pass
                    
                    with open(fname, 'wb') as f:
                        f.write(fh.getvalue())
                    
                    ext_type = os.path.splitext(fname)[1].lower().strip('.')
                    if not ext_type: 
                        ext_type = 'pdf'
                        
                    text = self.extractor.extract_text(fname, ext_type)
                    
                    try:
                        os.remove(fname)
                    except PermissionError:
                        pass
                    
                    if text:
                        yield {"text": text, "metadata": row}, idx, len(records)
                        
                except Exception as e:
                    logger.error(f"Row {idx} error: {e}")
        except Exception as e:
            logger.error(f"Sheet access error: {e}")
            yield {"error": str(e)}, 0, 0

class ResumeAnalyzer:
    def __init__(self, api_key: str):
        self.client = Groq(api_key=api_key)
    
    def analyze_resume(self, text: str, meta: Dict = None) -> Dict:
        prompt = f"""Analyze this resume and extract data as JSON.

Resume Text:
{text[:CONFIG['max_resume_chars']]}

JSON format:
{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "1234567890",
    "skills": ["Skill1", "Skill2", "Skill3"],
    "education": "Degree - University",
    "experience_years": 0.0
}}

RULES:
1. **Experience Calculation**: 
   - Identify every job/role and its duration.
   - Sum the total duration in years.
   - If a job is "Present" or "Current", assume end date is 2025.
   - Example: 2018-2020 (2y) + 2021-Present (4y) = 6.0 years.
   - Return strictly a number (float).
2. **Skills**: Extract technical, tools, and soft skills (Min 10).
3. **Phone**: Extract only digits (10-15 digits).
4. **Name**: If not found in text, use filename from metadata.
"""
        for attempt in range(CONFIG["max_retries"]):
            try:
                resp = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model=CONFIG["groq_model"],
                    temperature=CONFIG["temperature"],
                    max_tokens=CONFIG["max_tokens"],
                    response_format={"type": "json_object"}
                )
                
                data = json.loads(resp.choices[0].message.content)
                
                if not data.get("phone") or data["phone"] == "N/A":
                    data["phone"] = self._extract_phone(text)
                else:
                    data["phone"] = re.sub(r'\D', '', str(data["phone"]))[-10:]
                
                if not data.get("email") or data["email"] == "N/A":
                    data["email"] = self._extract_regex(text, r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
                
                if not data.get("skills") or len(data.get("skills", [])) < 5:
                    extracted = self._extract_skills(text)
                    if len(extracted) > len(data.get("skills", [])):
                        data["skills"] = extracted
                
                if not data.get("experience_years") or data.get("experience_years") == 0:
                    data["experience_years"] = self._calc_exp_fallback(text)
                
                data["ats_score"] = self._calc_score_enhanced(data)
                
                if meta:
                    if data.get("name") == "Full Name" or not data.get("name"):
                        for k, v in meta.items():
                            if "name" in k.lower():
                                data["name"] = v
                                break

                return data
                
            except Exception as e:
                if attempt < CONFIG["max_retries"] - 1:
                    time.sleep(CONFIG["retry_delay"] * (attempt + 1))
                else:
                    return self._create_fallback(text, meta)
        
        return self._create_fallback(text, meta)
    
    def _extract_phone(self, text: str) -> str:
        patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\b\d{10}\b'
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text[:2000])
            if matches:
                digits = re.sub(r'\D', '', matches[0])
                if len(digits) >= 10:
                    return digits[-10:]
        return "N/A"
    
    def _extract_regex(self, text: str, pattern: str) -> str:
        m = re.findall(pattern, text[:2000])
        return (''.join(m[0]) if isinstance(m[0], tuple) else m[0]) if m else "N/A"
    
    def _extract_skills(self, text: str) -> List[str]:
        skill_keywords = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'swift', 'kotlin', 'go',
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring',
            'sql', 'mysql', 'postgresql', 'mongodb', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
            'git', 'jenkins', 'ci/cd', 'linux', 'machine learning', 'tensorflow', 'pytorch', 'pandas',
            'agile', 'scrum', 'leadership', 'communication', 'problem solving'
        ]
        text_lower = text.lower()
        return [s.capitalize() for s in skill_keywords if re.search(r'\b' + re.escape(s) + r'\b', text_lower)]
    
    def _calc_exp_fallback(self, text: str) -> float:
        try:
            years_mentioned = re.findall(r'\b(19[9]\d|20[0-2]\d)\b', text)
            if len(years_mentioned) >= 2:
                years_list = sorted([int(y) for y in set(years_mentioned)])
                span = years_list[-1] - years_list[0]
                if 1 <= span <= 40:
                    return float(span)
            return 0.0
        except:
            return 0.0
    
    def _calc_score_enhanced(self, d: Dict) -> int:
        score = 0
        if d.get("name") and d["name"] != "N/A": score += 10
        if d.get("email"): score += 10
        if d.get("phone") and d["phone"] != "N/A": score += 5
        score += min(len(d.get("skills", [])) * 2, 40)
        score += min(float(d.get("experience_years", 0)) * 5, 20)
        if d.get("education") and d["education"] != "N/A": score += 15
        return min(100, score)
    
    def _create_fallback(self, text: str, meta: Dict) -> Dict:
        return {
            "name": meta.get("Name", "Unknown") if meta else "Unknown",
            "email": self._extract_regex(text, r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            "phone": self._extract_phone(text),
            "skills": self._extract_skills(text),
            "education": "N/A",
            "experience_years": self._calc_exp_fallback(text),
            "ats_score": 20
        }

class RAGPipeline:
    def __init__(self, api_key: str):
        self.embeddings = HuggingFaceEmbeddings(
            model_name=CONFIG["embedding_model"],
            model_kwargs={'device': 'cpu'}
        )
        self.llm = ChatGroq(
            groq_api_key=api_key,
            model_name=CONFIG["groq_model"],
            temperature=0.1,
            max_tokens=800
        )
        self.candidates_data = []
    
    def build_vector_store(self, candidates: List[Dict]):
        self.candidates_data = candidates
        docs = []
        
        for idx, c in enumerate(candidates):
            skills_str = ', '.join([str(s) for s in c.get('skills', [])])
            
            content = f"""CANDIDATE #{idx+1}
NAME: {c.get('name', 'Unknown')}
SCORE: {c.get('ats_score', 0)}
EXP: {c.get('experience_years', 0)} years
SKILLS: {skills_str}
EDUCATION: {c.get('education', 'N/A')}
"""
            docs.append(LangchainDocument(
                page_content=content,
                metadata={"candidate_id": idx, "name": c.get('name'), "ats_score": c.get('ats_score', 0)}
            ))
        
        splits = RecursiveCharacterTextSplitter(chunk_size=CONFIG["chunk_size"], chunk_overlap=0).split_documents(docs)
        self.vector_store = FAISS.from_documents(splits, self.embeddings)
        
        template = """Context: {context}
Question: {question}
Answer concisely based on context."""
        
        self.qa_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm,
            retriever=self.vector_store.as_retriever(search_kwargs={"k": 5}),
            return_source_documents=True,
            combine_docs_chain_kwargs={"prompt": PromptTemplate(template=template, input_variables=["context", "question"])}
        )
    
    def query(self, q: str) -> Dict:
        try:
            result = self.qa_chain({"question": q, "chat_history": []})
            sources = list({s['candidate_id']: s for s in [d.metadata for d in result.get("source_documents", [])]}.values())[:3]
            return {"answer": result["answer"], "sources": sources}
        except Exception as e:
            return {"error": str(e)}
    
    def get_candidate(self, idx: int) -> Optional[Dict]:
        return self.candidates_data[idx] if 0 <= idx < len(self.candidates_data) else None

def load_api_key() -> str:
    return json.load(open(CONFIG["config_file"])).get("groq_api_key", "") if os.path.exists(CONFIG["config_file"]) else ""

def save_api_key(k: str):
    with open(CONFIG["config_file"], 'w') as f:
        json.dump({"groq_api_key": k}, f)

def process_resume_parallel(args):
    text, meta, analyzer = args
    return analyzer.analyze_resume(text, meta)

def main():
    st.set_page_config(page_title="Resume Analyzer", page_icon="📋", layout="wide")
    
    st.markdown("""<style>
    .stChatFloatingInputContainer { position: fixed; bottom: 0; left: 0; right: 0; background: white; padding: 1rem; z-index: 999; }
    .main .block-container { padding-bottom: 120px; }
    .stButton button { border-radius: 8px; }
    </style>""", unsafe_allow_html=True)
    
    for k in ["candidates", "rag_pipeline", "chat_history", "trigger_query"]:
        if k not in st.session_state:
            st.session_state[k] = [] if k != "rag_pipeline" and k != "trigger_query" else None
    
    if "trigger_query" not in st.session_state:
        st.session_state.trigger_query = False

    api_key = load_api_key()
    
    if not api_key:
        st.title("🔐 Setup Required")
        with st.form("api"):
            k = st.text_input("Groq API Key", type="password", placeholder="gsk_...")
            if st.form_submit_button("Save"):
                save_api_key(k)
                st.rerun()
        st.stop()
    
    st.title("📋 Resume Analyzer & Chatbot")
    
    with st.sidebar:
        st.markdown("### ⚙️ Settings")
        if st.button("🔄 Reset API"):
            os.remove(CONFIG["config_file"])
            st.rerun()
        if st.button("🗑️ Clear Data"):
            st.session_state.clear()
            st.rerun()
        
        st.markdown("---")
        if st.session_state.candidates:
            st.metric("Total Candidates", len(st.session_state.candidates))
    
    tab1, tab2, tab3 = st.tabs(["📥 Upload", "💬 Chatbot", "📊 Results"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            files = st.file_uploader("Upload Files", type=['pdf', 'docx', 'png', 'jpg'], accept_multiple_files=True)
        with col2:
            sheet = st.text_input("Google Sheet URL")
        
        if st.button("🚀 Process", type="primary"):
            if not files and not sheet:
                st.error("No data provided")
                st.stop()
            
            start = datetime.now()
            ext = DocumentExtractor()
            analyzer = ResumeAnalyzer(api_key)
            candidates = []
            file_data = []
            
            if files:
                prog = st.progress(0)
                status = st.empty()
                for idx, f in enumerate(files):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{f.name.split('.')[-1]}") as tmp:
                        tmp.write(f.getvalue())
                        tmp_path = tmp.name
                    
                    text = ext.extract_text(tmp_path, f.name.split('.')[-1].lower())
                    try: os.remove(tmp_path)
                    except: pass
                    
                    if text:
                        file_data.append((text, {"filename": f.name}, analyzer))
                    prog.progress((idx + 1) / len(files))
                    status.text(f"Extracted {idx+1}/{len(files)}")
            
            if file_data:
                batch_size = 3
                for i in range(0, len(file_data), batch_size):
                    batch = file_data[i:i+batch_size]
                    with ThreadPoolExecutor(max_workers=CONFIG["max_workers"]) as exe:
                        futures = [exe.submit(process_resume_parallel, x) for x in batch]
                        for f in as_completed(futures):
                            if f.result(): candidates.append(f.result())
                    time.sleep(1)
            
            if sheet:
                handler = GoogleDriveHandler(ext)
                for data, _, _ in handler.process_sheet(sheet):
                    if "error" in data:
                        st.error(data["error"])
                        continue
                    if data.get("text"):
                        candidates.append(analyzer.analyze_resume(data["text"], data.get("metadata")))

            if candidates:
                st.session_state.candidates = sorted(candidates, key=lambda x: x.get("ats_score", 0), reverse=True)
                rag = RAGPipeline(api_key)
                rag.build_vector_store(st.session_state.candidates)
                st.session_state.rag_pipeline = rag
                st.success(f"Processed {len(candidates)} resumes!")
                st.balloons()
            else:
                st.error("No data extracted.")

    with tab2:
        if not st.session_state.rag_pipeline:
            st.info("Upload resumes first.")
        else:
            qs = ["Top candidate?", "Python developers?", "Exp > 5 years?"]
            cols = st.columns(3)
            for i, q in enumerate(qs):
                if cols[i].button(q, use_container_width=True):
                    st.session_state.chat_history.append({"role": "user", "content": q})
                    st.session_state.trigger_query = True
            
            if prompt := st.chat_input("Ask question..."):
                st.session_state.chat_history.append({"role": "user", "content": prompt})
                st.session_state.trigger_query = True
            
            if st.session_state.trigger_query and st.session_state.chat_history:
                last_msg = st.session_state.chat_history[-1]
                if last_msg["role"] == "user":
                    with st.chat_message("user"):
                        st.markdown(last_msg["content"])
                        
                    with st.chat_message("assistant"):
                        with st.spinner("Thinking..."):
                            res = st.session_state.rag_pipeline.query(last_msg["content"])
                            if "error" in res:
                                ans = f"Error: {res['error']}"
                            else:
                                ans = res["answer"]
                                if res.get("sources"):
                                    ans += "\n\n**Sources:**\n" + "\n".join([f"- {s.get('name')} ({s.get('ats_score')})" for s in res["sources"]])
                            st.markdown(ans)
                            st.session_state.chat_history.append({"role": "assistant", "content": ans})
                
                st.session_state.trigger_query = False
                st.rerun()

            for msg in reversed(st.session_state.chat_history):
                if msg == st.session_state.chat_history[-1] and st.session_state.trigger_query:
                    continue
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

    with tab3:
        if st.session_state.candidates:
            # Modified Column Sequence
            filtered_data = []
            columns_to_keep = ['name', 'education', 'ats_score', 'skills', 'email', 'phone', 'experience_years']
            
            for c in st.session_state.candidates:
                row = {k: c.get(k, 'N/A') for k in columns_to_keep}
                
                if isinstance(row['skills'], list):
                    row['skills'] = ", ".join([str(s) for s in row['skills']])
                
                row['phone'] = re.sub(r'\D', '', str(row['phone']))
                
                filtered_data.append(row)
            
            df = pd.DataFrame(filtered_data)
            
            # Rename for display
            display_df = df.rename(columns={
                'name': 'Name', 'education': 'Education', 'ats_score': 'ATS Score', 
                'skills': 'Skills', 'email': 'Email', 'phone': 'Phone', 
                'experience_years': 'Exp (Years)'
            })
            
            st.dataframe(display_df, use_container_width=True)
            
            csv_buf = io.StringIO()
            csv_buf.write(','.join(display_df.columns) + '\n')
            
            for _, row in display_df.iterrows():
                row_dict = row.to_dict()
                phone_val = row_dict['Phone']
                if phone_val and str(phone_val) != 'N/A':
                    row_dict['Phone'] = f'="{phone_val}"' 
                
                row_vals = [f'"{str(v)}"' if k != 'Phone' else str(v) for k, v in row_dict.items()]
                csv_buf.write(','.join(row_vals) + '\n')
            
            st.download_button("Download CSV", csv_buf.getvalue(), "resumes.csv", "text/csv")

if __name__ == "__main__":
    main()