"""
ATS Resume Analyzer - Production Ready | Enhanced Version
"""

import streamlit as st
import json, tempfile, os, re, io, logging, sys, time
from datetime import datetime
from typing import List, Dict, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

# Windows-compatible logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('ats.log', mode='w', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

try:
    import fitz  # PyMuPDF
    import pandas as pd
    from rapidocr_onnxruntime import RapidOCR
    from docx import Document
    from groq import Groq
    import gspread
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain.docstore.document import Document as LangchainDocument
    from langchain.chains import ConversationalRetrievalChain
    from langchain_groq import ChatGroq
    from langchain.memory import ConversationBufferMemory
    from langchain.prompts import PromptTemplate
    logger.info("All dependencies loaded successfully")
except ImportError as e:
    st.error(f"Missing library: {e}\nRun: pip install -r requirements.txt")
    st.stop()

CONFIG = {
    "groq_model": "llama-3.1-8b-instant",
    "embedding_model": "sentence-transformers/all-MiniLM-L6-v2",
    "chunk_size": 4000,  
    "chunk_overlap": 0,
    "max_resume_chars": 8000,
    "max_tokens": 1500,
    "temperature": 0.1,
    "max_workers": 4,
    "retry_delay": 2,
    "max_retries": 3
}

GOOGLE_SCOPES = [
    'https://www.googleapis.com/auth/spreadsheets',
    'https://www.googleapis.com/auth/drive.readonly'
]

class DocumentExtractor:
    def __init__(self):
        self.ocr = RapidOCR()
    
    def extract_text(self, path: str, ftype: str) -> str:
        try:
            if ftype == 'pdf':
                text_content = ""
                with fitz.open(path) as doc:
                    for page in doc:
                        page_text = page.get_text("text").strip()
                        if len(page_text) > 50: 
                            text_content += page_text + "\n"
                        else:
                            try:
                                pix = page.get_pixmap()
                                img_data = pix.tobytes("png")
                                result, _ = self.ocr(img_data)
                                if result:
                                    text_content += "\n".join([line[1] for line in result]) + "\n"
                            except Exception as e:
                                logger.error(f"OCR failed on PDF page: {e}")
                return text_content.strip()
            elif ftype == 'docx':
                return "\n".join([p.text.strip() for p in Document(path).paragraphs if p.text.strip()])
            elif ftype in ['png', 'jpg', 'jpeg']:
                result, _ = self.ocr(path)
                return "\n".join([line[1] for line in result]) if result else ""
        except Exception as e:
            logger.error(f"Extract error: {e}")
        return ""

class GoogleDriveHandler:
    def __init__(self, extractor):
        self.extractor = extractor
        creds = None
        
        # Try loading from streamlit secrets first (recommended for Cloud)
        if "google_credentials" in st.secrets:
            creds = Credentials.from_authorized_user_info(st.secrets["google_credentials"], GOOGLE_SCOPES)
        elif os.path.exists('token.json'):
            creds = Credentials.from_authorized_user_file('token.json', GOOGLE_SCOPES)
            
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    creds.refresh(Request())
                except:
                    creds = None
            
            if not creds:
                # On Streamlit Cloud, we can't run a local server
                if os.environ.get("STREAMLIT_RUNTIME_ENV") == "cloud" or not os.path.exists('credentials.json'):
                    logger.warning("Google Credentials missing or invalid. Skipping Drive integration.")
                    self.gc = None
                    self.drive = None
                    return
                
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', GOOGLE_SCOPES)
                creds = flow.run_local_server(port=0)
                with open('token.json', 'w') as token:
                    token.write(creds.to_json())
                    
        self.gc = gspread.authorize(creds)
        self.drive = build('drive', 'v3', credentials=creds)
    
    def process_sheet(self, url: str):
        if not self.gc:
            yield {"error": "Google Drive integration not configured. Please use direct file upload or check documentation for setup."}, 0, 0
            return
        try:
            sheet = self.gc.open_by_url(url).get_worksheet(0)
            records = sheet.get_all_records()
            
            link_col = next((k for r in records[:5] for k, v in r.items() if isinstance(v, str) and "drive.google.com" in v), None)
            
            if not link_col:
                yield {"error": "No Drive links found"}, 0, 0
                return
            
            for idx, row in enumerate(records, 1):
                try:
                    link = row.get(link_col, "").strip()
                    if not link:
                        continue
                    
                    match = re.search(r'/file/d/([a-zA-Z0-9_-]+)|id=([a-zA-Z0-9_-]+)', link)
                    if not match:
                        continue
                    fid = match.group(1) or match.group(2)
                    
                    meta = self.drive.files().get(fileId=fid).execute()
                    fname = f"temp_{fid}_{meta.get('name', 'file')}"
                    
                    request = self.drive.files().get_media(fileId=fid)
                    fh = io.BytesIO()
                    downloader = MediaIoBaseDownload(fh, request)
                    while not downloader.next_chunk()[1]:
                        pass
                    
                    with open(fname, 'wb') as f:
                        f.write(fh.getvalue())
                    
                    ext_type = os.path.splitext(fname)[1].lower().strip('.')
                    if not ext_type: 
                        ext_type = 'pdf'
                        
                    text = self.extractor.extract_text(fname, ext_type)
                    
                    try:
                        os.remove(fname)
                    except PermissionError:
                        pass
                    
                    if text:
                        yield {"text": text, "metadata": row}, idx, len(records)
                        
                except Exception as e:
                    logger.error(f"Row {idx} error: {e}")
        except Exception as e:
            logger.error(f"Sheet access error: {e}")
            yield {"error": str(e)}, 0, 0

class ResumeAnalyzer:
    def __init__(self, api_key: str):
        self.client = Groq(api_key=api_key)
    
    def analyze_resume(self, text: str, meta: Dict = None, jd: str = None) -> Dict:
        jd_prompt = f"\n\nTarget Job Description:\n{jd}" if jd else ""
        prompt = f"""Analyze this resume and extract data as JSON. 
If a Job Description is provided, calculate a match score based on skills, experience, and role alignment.{jd_prompt}

Resume Text:
{text[:CONFIG['max_resume_chars']]}

JSON format:
{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "1234567890",
    "skills": ["Skill1", "Skill2"],
    "education": "Degree - University",
    "experience_years": 0.0,
    "match_percentage": 0,
    "matching_skills": ["S1"],
    "missing_skills": ["S2"],
    "summary": "Brief 1-sentence professional summary",
    "certifications": ["Cert1"],
    "education_level": "Bachelors/Masters/PhD"
}}

RULES:
1. **Experience**: Be precise. Sum total years across all roles. 
2. **Match Score**: Be critical. A 100% match should be extremely rare. Compare specific technical depth.
3. **Skills**: Categorize into Technical, Soft, and Tools if possible (return as flat list).
4. **Summary**: Write a compelling one-sentence pitch for this candidate.
"""
        for attempt in range(CONFIG["max_retries"]):
            try:
                resp = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model=CONFIG["groq_model"],
                    temperature=CONFIG["temperature"],
                    max_tokens=CONFIG["max_tokens"],
                    response_format={"type": "json_object"}
                )
                
                data = json.loads(resp.choices[0].message.content)
                
                # Post-processing
                data["phone"] = self._extract_phone(text) if not data.get("phone") else re.sub(r'\D', '', str(data["phone"]))[-10:]
                data["email"] = self._extract_regex(text, r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b') if not data.get("email") else data["email"]
                
                if not data.get("experience_years"):
                    data["experience_years"] = self._calc_exp_fallback(text)
                
                # ATS Score is a composite of data quality + JD match
                base_score = self._calc_score_enhanced(data)
                if jd and data.get("match_percentage", 0) > 0:
                    data["ats_score"] = int((base_score * 0.3) + (data["match_percentage"] * 0.7))
                else:
                    data["ats_score"] = base_score

                if meta:
                    if data.get("name") == "Full Name" or not data.get("name"):
                        data["name"] = meta.get("filename", meta.get("Name", "Unknown"))

                return data
                
            except Exception as e:
                if attempt < CONFIG["max_retries"] - 1:
                    time.sleep(CONFIG["retry_delay"] * (attempt + 1))
                else:
                    return self._create_fallback(text, meta)
        
        return self._create_fallback(text, meta)
    
    def _extract_phone(self, text: str) -> str:
        patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\b\d{10}\b'
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text[:2000])
            if matches:
                digits = re.sub(r'\D', '', matches[0])
                if len(digits) >= 10:
                    return digits[-10:]
        return "N/A"
    
    def _extract_regex(self, text: str, pattern: str) -> str:
        m = re.findall(pattern, text[:2000])
        return (''.join(m[0]) if isinstance(m[0], tuple) else m[0]) if m else "N/A"
    
    def _extract_skills(self, text: str) -> List[str]:
        skill_keywords = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'swift', 'kotlin', 'go',
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring',
            'sql', 'mysql', 'postgresql', 'mongodb', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
            'git', 'jenkins', 'ci/cd', 'linux', 'machine learning', 'tensorflow', 'pytorch', 'pandas',
            'agile', 'scrum', 'leadership', 'communication', 'problem solving'
        ]
        text_lower = text.lower()
        return [s.capitalize() for s in skill_keywords if re.search(r'\b' + re.escape(s) + r'\b', text_lower)]
    
    def _calc_exp_fallback(self, text: str) -> float:
        try:
            years_mentioned = re.findall(r'\b(19[9]\d|20[0-2]\d)\b', text)
            if len(years_mentioned) >= 2:
                years_list = sorted([int(y) for y in set(years_mentioned)])
                span = years_list[-1] - years_list[0]
                if 1 <= span <= 40:
                    return float(span)
            return 0.0
        except:
            return 0.0
    
    def _calc_score_enhanced(self, d: Dict) -> int:
        score = 0
        if d.get("name") and d["name"] != "N/A": score += 10
        if d.get("email"): score += 10
        if d.get("phone") and d["phone"] != "N/A": score += 5
        score += min(len(d.get("skills", [])) * 2, 40)
        score += min(float(d.get("experience_years", 0)) * 5, 20)
        if d.get("education") and d["education"] != "N/A": score += 15
        return min(100, score)
    
    def _create_fallback(self, text: str, meta: Dict) -> Dict:
        return {
            "name": meta.get("Name", "Unknown") if meta else "Unknown",
            "email": self._extract_regex(text, r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            "phone": self._extract_phone(text),
            "skills": self._extract_skills(text),
            "education": "N/A",
            "experience_years": self._calc_exp_fallback(text),
            "ats_score": 20
        }

class RAGPipeline:
    def __init__(self, api_key: str):
        self.embeddings = HuggingFaceEmbeddings(
            model_name=CONFIG["embedding_model"],
            model_kwargs={'device': 'cpu'}
        )
        self.llm = ChatGroq(
            groq_api_key=api_key,
            model_name=CONFIG["groq_model"],
            temperature=0.1,
            max_tokens=800
        )
        self.candidates_data = []
    
    def build_vector_store(self, candidates: List[Dict]):
        self.candidates_data = candidates
        docs = []
        
        for idx, c in enumerate(candidates):
            skills_str = ', '.join([str(s) for s in c.get('skills', [])])
            
            content = f"""CANDIDATE #{idx+1}
NAME: {c.get('name', 'Unknown')}
SCORE: {c.get('ats_score', 0)}
EXP: {c.get('experience_years', 0)} years
SKILLS: {skills_str}
EDUCATION: {c.get('education', 'N/A')}
"""
            docs.append(LangchainDocument(
                page_content=content,
                metadata={"candidate_id": idx, "name": c.get('name'), "ats_score": c.get('ats_score', 0)}
            ))
        
        splits = RecursiveCharacterTextSplitter(chunk_size=CONFIG["chunk_size"], chunk_overlap=0).split_documents(docs)
        self.vector_store = FAISS.from_documents(splits, self.embeddings)
        
        template = """Context: {context}
Question: {question}
Answer concisely based on context."""
        
        self.qa_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm,
            retriever=self.vector_store.as_retriever(search_kwargs={"k": 5}),
            return_source_documents=True,
            combine_docs_chain_kwargs={"prompt": PromptTemplate(template=template, input_variables=["context", "question"])}
        )
    
    def query(self, q: str) -> Dict:
        try:
            result = self.qa_chain({"question": q, "chat_history": []})
            sources = list({s['candidate_id']: s for s in [d.metadata for d in result.get("source_documents", [])]}.values())[:3]
            return {"answer": result["answer"], "sources": sources}
        except Exception as e:
            return {"error": str(e)}
    
    def get_candidate(self, idx: int) -> Optional[Dict]:
        return self.candidates_data[idx] if 0 <= idx < len(self.candidates_data) else None

def get_api_key():
    """Retrieve API key from secrets, environment, or session state."""
    if "groq_api_key" in st.secrets:
        return st.secrets["groq_api_key"]
    if os.environ.get("GROQ_API_KEY"):
        return os.environ.get("GROQ_API_KEY")
    return st.session_state.get("custom_api_key", "")

def save_api_key(k: str):
    st.session_state["custom_api_key"] = k

def process_resume_parallel(args):
    text, meta, analyzer, jd = args
    return analyzer.analyze_resume(text, meta, jd)

def apply_custom_style():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    html, body, [data-testid="stSidebar"] {
        font-family: 'Inter', sans-serif;
    }
    
    .main {
        background-color: #f8fafc;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
        background-color: transparent;
    }

    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: white;
        border-radius: 10px 10px 0px 0px;
        gap: 1px;
        padding-top: 10px;
        padding-bottom: 10px;
        border: 1px solid #e2e8f0;
    }

    .stTabs [aria-selected="true"] {
        background-color: #3b82f6 !important;
        color: white !important;
    }
    
    div[data-testid="stMetricValue"] {
        font-size: 2rem;
        font-weight: 700;
        color: #1e293b;
    }
    
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        background-color: #3b82f6;
        color: white;
        font-weight: 600;
        border: none;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        background-color: #2563eb;
        transform: translateY(-2px);
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
    }
    
    .css-1offfwp {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 2rem;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    
    /* Chat styling */
    .stChatMessage {
        border-radius: 12px;
        padding: 1rem;
        margin-bottom: 1rem;
    }
    /* Dataframe styling */
    .stDataFrame {
        border: 1px solid #e2e8f0;
        border-radius: 8px;
    }
    
    /* Metrics styling */
    [data-testid="stMetric"] {
        background-color: white;
        padding: 1rem;
        border-radius: 12px;
        border: 1px solid #e2e8f0;
        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1);
    }
    </style>
    """, unsafe_allow_html=True)

def main():
    st.set_page_config(page_title="ATS Resume Analyzer Pro", page_icon="🚀", layout="wide")
    apply_custom_style()
    
    # Initialize session state
    for k in ["candidates", "rag_pipeline", "chat_history", "trigger_query", "jd_text"]:
        if k not in st.session_state:
            st.session_state[k] = [] if k in ["candidates", "chat_history"] else (None if k != "jd_text" else "")
    
    if "trigger_query" not in st.session_state:
        st.session_state.trigger_query = False

    api_key = get_api_key()
    
    with st.sidebar:
        st.title("⚙️ Configuration")
        if not api_key:
            st.warning("Groq API Key not found!")
            new_key = st.text_input("Enter Groq API Key", type="password")
            if st.button("Save Key"):
                save_api_key(new_key)
                st.rerun()
        else:
            st.success("API Key Active")
            if st.button("Change API Key"):
                st.session_state.custom_api_key = ""
                st.rerun()
        
        st.markdown("---")
        st.markdown("### 📋 Job Description")
        st.session_state.jd_text = st.text_area("Paste the target Job Description here to get accuracy scores.", 
                                               value=st.session_state.jd_text, height=250)
        
        st.markdown("---")
        with st.expander("📖 How to Use"):
            st.markdown("""
            1. **Set API Key**: Enter your Groq API key above.
            2. **Add JD**: Paste the Job Description in the text area.
            3. **Upload**: Drag & drop resumes (PDF, Docx, Images).
            4. **Analyze**: Click '🚀 Process' to start.
            5. **Insights**:
               - **Chatbot**: Ask questions about candidates.
               - **Results**: See scores, matches, and visualizations.
            """)
        
        if st.button("🗑️ Clear All Data"):
            st.session_state.clear()
            st.rerun()
    
    tab1, tab2, tab3 = st.tabs(["📥 Upload", "💬 Chatbot", "📊 Results"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            files = st.file_uploader("Upload Files", type=['pdf', 'docx', 'png', 'jpg'], accept_multiple_files=True)
        with col2:
            sheet = st.text_input("Google Sheet URL")
        
        if st.button("🚀 Process", type="primary"):
            if not files and not sheet:
                st.error("No data provided")
                st.stop()
            
            start = datetime.now()
            ext = DocumentExtractor()
            analyzer = ResumeAnalyzer(api_key)
            candidates = []
            file_data = []
            if files:
                prog = st.progress(0)
                status = st.empty()
                for idx, f in enumerate(files):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{f.name.split('.')[-1]}") as tmp:
                        tmp.write(f.getvalue())
                        tmp_path = tmp.name
                    
                    text = ext.extract_text(tmp_path, f.name.split('.')[-1].lower())
                    try: os.remove(tmp_path)
                    except: pass
                    
                    if text:
                        file_data.append((text, {"filename": f.name}, analyzer, st.session_state.jd_text))
                    prog.progress((idx + 1) / len(files))
                    status.text(f"Extracted {idx+1}/{len(files)}")
            
            if file_data:
                with st.spinner("Analyzing resumes with AI..."):
                    batch_size = 4
                    for i in range(0, len(file_data), batch_size):
                        batch = file_data[i:i+batch_size]
                        with ThreadPoolExecutor(max_workers=CONFIG["max_workers"]) as exe:
                            futures = [exe.submit(process_resume_parallel, x) for x in batch]
                            for f in as_completed(futures):
                                res = f.result()
                                if res: candidates.append(res)
                        time.sleep(0.5)
            
            if sheet:
                handler = GoogleDriveHandler(ext)
                for data, _, _ in handler.process_sheet(sheet):
                    if "error" in data:
                        st.error(data["error"])
                        continue
                    if data.get("text"):
                        candidates.append(analyzer.analyze_resume(data["text"], data.get("metadata"), st.session_state.jd_text))

            if candidates:
                st.session_state.candidates = sorted(candidates, key=lambda x: x.get("ats_score", 0), reverse=True)
                rag = RAGPipeline(api_key)
                rag.build_vector_store(st.session_state.candidates)
                st.session_state.rag_pipeline = rag
                st.success(f"Successfully processed {len(candidates)} resumes!")
                st.balloons()
            else:
                st.error("No data could be extracted. Please check file formats.")

    with tab2:
        if not st.session_state.rag_pipeline:
            st.info("Upload resumes first.")
        else:
            qs = ["Top candidate?", "Python developers?", "Exp > 5 years?"]
            cols = st.columns(3)
            for i, q in enumerate(qs):
                if cols[i].button(q, use_container_width=True):
                    st.session_state.chat_history.append({"role": "user", "content": q})
                    st.session_state.trigger_query = True
            
            if prompt := st.chat_input("Ask question..."):
                st.session_state.chat_history.append({"role": "user", "content": prompt})
                st.session_state.trigger_query = True
            
            if st.session_state.trigger_query and st.session_state.chat_history:
                last_msg = st.session_state.chat_history[-1]
                if last_msg["role"] == "user":
                    with st.chat_message("user"):
                        st.markdown(last_msg["content"])
                        
                    with st.chat_message("assistant"):
                        with st.spinner("Thinking..."):
                            res = st.session_state.rag_pipeline.query(last_msg["content"])
                            if "error" in res:
                                ans = f"Error: {res['error']}"
                            else:
                                ans = res["answer"]
                                if res.get("sources"):
                                    ans += "\n\n**Sources:**\n" + "\n".join([f"- {s.get('name')} ({s.get('ats_score')})" for s in res["sources"]])
                            st.markdown(ans)
                            st.session_state.chat_history.append({"role": "assistant", "content": ans})
                
                st.session_state.trigger_query = False
                st.rerun()

            for msg in reversed(st.session_state.chat_history):
                if msg == st.session_state.chat_history[-1] and st.session_state.trigger_query:
                    continue
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

    with tab3:
        if st.session_state.candidates:
            st.header("📊 Analysis Overview")
            
            # Overview Metrics
            m1, m2, m3 = st.columns(3)
            avg_score = sum(c.get('ats_score', 0) for c in st.session_state.candidates) / len(st.session_state.candidates)
            top_score = max(c.get('ats_score', 0) for c in st.session_state.candidates)
            m1.metric("Avg ATS Score", f"{avg_score:.1f}")
            m2.metric("Top Score", f"{top_score}")
            m3.metric("Candidates", len(st.session_state.candidates))

            # Distribution Chart
            import plotly.express as px
            df_scores = pd.DataFrame([{"Name": c.get('name'), "Score": c.get('ats_score', 0)} for c in st.session_state.candidates])
            fig = px.bar(df_scores, x="Name", y="Score", title="ATS Score Distribution", 
                         color="Score", color_continuous_scale="RdYlGn")
            st.plotly_chart(fig, use_container_width=True)

            st.divider()
            
            # Main Data Table
            filtered_data = []
            columns_to_keep = ['name', 'ats_score', 'match_percentage', 'experience_years', 'skills', 'education', 'email']
            for c in st.session_state.candidates:
                row = {k: c.get(k, 'N/A') for k in columns_to_keep}
                if isinstance(row['skills'], list):
                    row['skills'] = ", ".join(row['skills'])
                filtered_data.append(row)
            
            df = pd.DataFrame(filtered_data)
            st.dataframe(df.rename(columns={
                'name': 'Name', 'ats_score': 'ATS Score', 'match_percentage': 'JD Match %',
                'experience_years': 'Exp (Yrs)', 'skills': 'Top Skills', 'education': 'Education'
            }), use_container_width=True)

            # Detailed Candidate View
            st.subheader("🔍 Detailed Candidate Insights")
            selected_name = st.selectbox("Select Candidate to view details", [c.get('name') for c in st.session_state.candidates])
            if selected_name:
                cand = next(c for c in st.session_state.candidates if c.get('name') == selected_name)
                c1, c2 = st.columns([1, 2])
                with c1:
                    st.metric("ATS Score", cand.get('ats_score', 0))
                    st.metric("JD Match", f"{cand.get('match_percentage', 0)}%")
                with c2:
                    st.write(f"**Summary:** {cand.get('summary', 'N/A')}")
                    st.write(f"**Matching Skills:** {', '.join(cand.get('matching_skills', []))}")
                    st.write(f"**Missing Skills:** {', '.join(cand.get('missing_skills', []))}")
            
            st.divider()
            csv_buf = io.StringIO()
            df.to_csv(csv_buf, index=False)
            st.download_button("📥 Download Full Report (CSV)", csv_buf.getvalue(), "ats_results.csv", "text/csv")
        else:
            st.info("No candidates processed yet. Go to the Upload tab to start.")

if __name__ == "__main__":
    main()